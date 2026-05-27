#!/usr/bin/env python3
"""
Meeting Notes DOCX Generator

Generates professionally formatted Word documents for meeting notes
with company branding:
- Font: Avenir Next For company (falls back to Avenir)
- Colors: Peppercorn (#241C15) headers, Cavendish Yellow (#FFE01B) accents

Usage:
    python generate_meeting_notes_docx.py <output_path>
"""

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)

# Branding
CAVENDISH_YELLOW = RGBColor(0xFF, 0xE0, 0x1B)
PEPPERCORN = RGBColor(0x24, 0x1C, 0x15)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BRAND_FONT = 'Avenir Next For company'


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def apply_font(paragraph, font_name=BRAND_FONT):
    """Apply font to all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name


def add_styled_heading(doc, text, level=1):
    """Add a heading with brand styling."""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = PEPPERCORN
        run.font.name = BRAND_FONT
    return heading


def generate_domain_auth_meeting_notes():
    """Generate meeting notes for Tx Domain Authentication review."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = BRAND_FONT
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Tx Domain Authentication – Entri Integration', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PEPPERCORN
        run.font.name = BRAND_FONT
    
    # Subtitle
    subtitle = doc.add_paragraph('Meeting Notes & Follow-up Questions')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.name = BRAND_FONT
    
    # Date
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font(date_para)
    
    doc.add_paragraph()
    
    # =========================================================================
    # KEY QUESTIONS SECTION (TOP PRIORITY)
    # =========================================================================
    
    questions_heading = doc.add_heading('KEY QUESTIONS – Needs Clarification', 1)
    for run in questions_heading.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Red for urgency
        run.font.name = BRAND_FONT
    
    questions = [
        ("API Key Storage", "What table on the messaging side is used to store the API key created during the integration setup?"),
        ("Unverified Root Domain", "If the root domain presented in the auth flow is not yet verified on the Acme Platform side, what happens? Is it possible for an unverified root domain to appear in this flow?"),
        ("Failed Auth Recovery", "After a failed Entri authentication, the only option shown is manual DNS record update. Is there an option to retry the Entri auth flow instead?"),
        ("Subdomain Flow – Missing Option", "When coming from the subdomain creation flow, there's no 'Use a different domain' option. Is this intentional, or should it be available?"),
        ("Status Flow Consistency", "Are the status states (failed auth, pending, verified, etc.) consistent across both the subdomain and root domain authentication flows?"),
        ("Restart Authentication", "While validating DNS records, what does the 'Restart Authentication' button actually do? Does it clear progress and start over, or retry verification?"),
    ]
    
    # Create questions table
    table = doc.add_table(rows=len(questions) + 1, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(['Topic', 'Question']):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, '241C15')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                run.font.name = BRAND_FONT
    
    # Question rows
    for idx, (topic, question) in enumerate(questions):
        row = table.rows[idx + 1]
        row.cells[0].text = topic
        row.cells[1].text = question
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = BRAND_FONT
    
    # Set column widths
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.0)
    
    doc.add_paragraph()
    
    # =========================================================================
    # MEETING OVERVIEW
    # =========================================================================
    
    add_styled_heading(doc, 'Meeting Overview', 1)
    
    overview = doc.add_paragraph()
    overview.add_run('Feature: ').bold = True
    overview.add_run('Tx Domain Authentication Using Entri (OBS Part 1)\n')
    overview.add_run('PM: ').bold = True
    overview.add_run('Andre Pardue\n')
    overview.add_run('Goal: ').bold = True
    overview.add_run('Move domain authentication into Acme Platform and use Entri to expedite the DNS setup process.')
    apply_font(overview)
    
    doc.add_paragraph()
    
    # =========================================================================
    # KEY POINTS COVERED
    # =========================================================================
    
    add_styled_heading(doc, 'Key Points Covered', 1)
    
    # API Key Flow
    add_styled_heading(doc, 'API Key Creation Flow', 2)
    
    api_points = [
        "API key and domain are stored on the messaging side (existing table, not new)",
        "Flow: Welcome page → Create API Key (with copy option) → Done → API call to messaging to save key",
        "Integration call flows from Acme Platform to messaging for key persistence",
        "API key created messaging confirms successful setup",
    ]
    for point in api_points:
        p = doc.add_paragraph(point, style='List Bullet')
        apply_font(p)
    
    # Domain Auth Flow
    add_styled_heading(doc, 'Domain Authentication Flow', 2)
    
    domain_points = [
        "Entire UI flow happens on Acme Platform side (not messaging)",
        "Option to use an entirely different domain from signup",
        "Option to use just the root domain",
        "Recommendation: Create subdomain for Tx sends (separate from marketing)",
        "Root domain from account signup is pre-populated in the flow",
        "Root domain is checked for verified status before proceeding",
    ]
    for point in domain_points:
        p = doc.add_paragraph(point, style='List Bullet')
        apply_font(p)
    
    # Subdomain Flow
    add_styled_heading(doc, 'Subdomain Authentication', 2)
    
    subdomain_points = [
        "Confirm auth flow sends information to messaging to create subdomain (marked as verified)",
        "'Start Authentication' button appears after subdomain is entered",
        "Authentication proceeds through Entri model",
    ]
    for point in subdomain_points:
        p = doc.add_paragraph(point, style='List Bullet')
        apply_font(p)
    
    # Manual Option
    add_styled_heading(doc, 'Manual Authentication Option', 2)
    
    p = doc.add_paragraph(
        "For users who don't want to provide credentials to Entri, a manual authentication option "
        "exists with self-guided DNS record setup instructions.",
        style='List Bullet'
    )
    apply_font(p)
    
    doc.add_paragraph()
    
    # =========================================================================
    # CONTEXT / WHY
    # =========================================================================
    
    add_styled_heading(doc, 'Strategic Context', 1)
    
    context = doc.add_paragraph()
    context.add_run('Problem: ').bold = True
    context.add_run(
        "84% of messaging Entry users and 70% of Non-messaging Entry users never start domain "
        "authentication after creating a Tx account. The current flow requires navigating to "
        "the messaging app via a hard-to-find 'Launch App' button.\n\n"
    )
    context.add_run('Solution: ').bold = True
    context.add_run(
        "Move domain authentication into Acme Platform and use Entri to streamline DNS setup, "
        "providing clearer onboarding checklist and reducing drop-off."
    )
    apply_font(context)
    
    # Save
    output_path = Path('/Users/glehman/Downloads/Tx_Domain_Auth_Meeting_Notes.docx')
    doc.save(str(output_path))
    print(f'✅ Created: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_domain_auth_meeting_notes()
