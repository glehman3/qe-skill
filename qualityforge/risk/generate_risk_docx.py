#!/usr/bin/env python3
"""
QualityForge DOCX Generator

Generates professionally formatted Word documents from QualityForge outputs
with proper Acme Platform/company branding:
- Font: Avenir Next For company (falls back to Avenir)
- Colors: Peppercorn (#241C15) headers, Cavendish Yellow (#FFE01B) accents

Supports both:
- Risk Analysis folders (contains risk_analysis_report.md)
- Test Jam folders (contains testjam_all_test_cases.csv)

Usage:
    python generate_risk_docx.py <folder_path>

Examples:
    # Risk analysis folder
    python generate_risk_docx.py test-jams/2026-02-13_tx-domain-auth-entri-risk-analysis
    
    # Test jam folder (coverage map only)
    python generate_risk_docx.py test-jams/2026-02-13_tx-domain-auth-entri

Outputs (depending on folder type):
    - risk_analysis_report.docx (if risk_analysis_report.md exists)
    - test_coverage_map.docx (if test_coverage_map.md exists)

Requirements:
    pip install python-docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)

# =============================================================================
# BRANDING CONSTANTS
# =============================================================================

# Acme Platform brand colors
CAVENDISH_YELLOW = RGBColor(0xFF, 0xE0, 0x1B)
PEPPERCORN = RGBColor(0x24, 0x1C, 0x15)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Severity colors
CRITICAL_RED = RGBColor(0xC0, 0x00, 0x00)
HIGH_ORANGE = RGBColor(0xED, 0x7D, 0x31)
MEDIUM_YELLOW = RGBColor(0xBF, 0x90, 0x00)
LOW_GREEN = RGBColor(0x53, 0x81, 0x35)

# company brand font
BRAND_FONT = 'Avenir Next For company'
FALLBACK_FONT = 'Avenir'


# =============================================================================
# DOCUMENT UTILITIES
# =============================================================================

def set_document_font(doc, font_name=BRAND_FONT):
    """Set the default font for the entire document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(11)
    
    # Set font for headings
    for i in range(10):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = font_name
        except KeyError:
            pass
    
    # Set font for list styles
    for style_name in ['List Bullet', 'List Number', 'List Paragraph']:
        try:
            list_style = doc.styles[style_name]
            list_style.font.name = font_name
        except KeyError:
            pass


def apply_font_to_paragraph(paragraph, font_name=BRAND_FONT):
    """Apply font to all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_with_header(doc, headers, rows, header_color="241C15", font_name=BRAND_FONT):
    """Create a table with styled header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                run.font.name = font_name
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = font_name
    
    return table


def add_styled_heading(doc, text, level=1):
    """Add a heading with brand styling."""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = PEPPERCORN
        run.font.name = BRAND_FONT
    return heading


# =============================================================================
# MARKDOWN PARSING UTILITIES
# =============================================================================

def parse_markdown_table(lines):
    """Parse a markdown table into headers and rows."""
    headers = []
    rows = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        
        # Parse cells
        cells = [c.strip() for c in line.split('|')[1:-1]]
        
        # Skip separator lines (contain only dashes)
        if all(set(c.replace('-', '').replace(':', '')) == set() or c.replace('-', '').replace(':', '') == '' for c in cells):
            continue
        
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    
    return headers, rows


def extract_section(content, section_name):
    """Extract content under a markdown section heading."""
    pattern = rf'^##\s+{re.escape(section_name)}\s*$'
    lines = content.split('\n')
    
    in_section = False
    section_lines = []
    
    for line in lines:
        if re.match(pattern, line, re.IGNORECASE):
            in_section = True
            continue
        elif in_section and re.match(r'^##\s+', line):
            break
        elif in_section:
            section_lines.append(line)
    
    return '\n'.join(section_lines).strip()


def parse_risk_blocks(content):
    """Parse RISK-### blocks from the markdown content."""
    risks = []
    
    # Match risk headers like "### RISK-001: Title [SEVERITY]" or "#### RISK-001: Title"
    pattern = r'####+\s*(RISK-\d+):\s*(.+?)(?:\s*\[(CRITICAL|HIGH|MEDIUM|LOW)\])?\s*$'
    
    lines = content.split('\n')
    current_risk = None
    
    for i, line in enumerate(lines):
        match = re.match(pattern, line, re.IGNORECASE)
        if match:
            if current_risk:
                risks.append(current_risk)
            
            current_risk = {
                'id': match.group(1),
                'title': match.group(2).strip(),
                'severity': match.group(3).upper() if match.group(3) else 'MEDIUM',
                'content': []
            }
        elif current_risk:
            # Check for next risk or section
            if re.match(r'^##\s+', line) and not line.startswith('###'):
                risks.append(current_risk)
                current_risk = None
            else:
                current_risk['content'].append(line)
    
    if current_risk:
        risks.append(current_risk)
    
    return risks


# =============================================================================
# DOCUMENT GENERATORS
# =============================================================================

def generate_risk_analysis_docx(folder_path, output_path=None):
    """
    Generate a formatted DOCX from risk_analysis_report.md.
    
    Args:
        folder_path: Path to the risk analysis folder
        output_path: Optional custom output path (defaults to folder_path/risk_analysis_report.docx)
    
    Returns:
        Path to the generated DOCX file
    """
    folder = Path(folder_path)
    md_path = folder / 'risk_analysis_report.md'
    
    if not md_path.exists():
        raise FileNotFoundError(f"Risk analysis report not found: {md_path}")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    set_document_font(doc, BRAND_FONT)
    
    # Extract title from first heading
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    title_text = title_match.group(1) if title_match else 'Quality Risk Analysis'
    
    # Title
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PEPPERCORN
        run.font.name = BRAND_FONT
    
    # Extract feature name from subtitle or folder name
    feature_name = folder.name.split('_', 1)[-1].replace('-', ' ').title() if '_' in folder.name else 'Feature Analysis'
    feature_name = feature_name.replace('-Risk-Analysis', '').replace('Risk Analysis', '').strip()
    
    # Info section
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Generated: ').bold = True
    info.add_run(f'{datetime.now().strftime("%Y-%m-%d")}\n')
    info.add_run('Source: ').bold = True
    info.add_run(str(md_path))
    apply_font_to_paragraph(info, BRAND_FONT)
    
    # Parse and render each major section
    sections = [
        ('Summary', 2),
        ('Unified Requirements Summary', 2),
        ('Codebase Touch Points', 2),
        ('Risks', 2),
        ('Open Questions', 2),
        ('Recommended Spikes', 2),
        ('Testing Complexity Assessment', 2),
        ('Assumptions', 2),
    ]
    
    for section_name, level in sections:
        section_content = extract_section(content, section_name)
        if section_content:
            add_styled_heading(doc, section_name, level)
            
            # Check for tables in section
            if '|' in section_content and '---' in section_content:
                # Find and parse tables
                table_lines = []
                non_table_lines = []
                in_table = False
                
                for line in section_content.split('\n'):
                    if line.strip().startswith('|'):
                        in_table = True
                        table_lines.append(line)
                    elif in_table and not line.strip():
                        # End of table
                        if table_lines:
                            headers, rows = parse_markdown_table(table_lines)
                            if headers and rows:
                                add_table_with_header(doc, headers, rows)
                                doc.add_paragraph()
                            table_lines = []
                        in_table = False
                    else:
                        if table_lines:
                            headers, rows = parse_markdown_table(table_lines)
                            if headers and rows:
                                add_table_with_header(doc, headers, rows)
                                doc.add_paragraph()
                            table_lines = []
                        in_table = False
                        non_table_lines.append(line)
                
                # Handle remaining table
                if table_lines:
                    headers, rows = parse_markdown_table(table_lines)
                    if headers and rows:
                        add_table_with_header(doc, headers, rows)
                
                # Add non-table content
                for line in non_table_lines:
                    if line.strip():
                        if line.strip().startswith('- '):
                            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif line.strip().startswith('* '):
                            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif re.match(r'^\d+\.\s', line.strip()):
                            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
                        else:
                            p = doc.add_paragraph(line.strip())
                        apply_font_to_paragraph(p, BRAND_FONT)
            else:
                # Plain text section
                for line in section_content.split('\n'):
                    if line.strip():
                        if line.strip().startswith('### '):
                            add_styled_heading(doc, line.strip()[4:], 3)
                        elif line.strip().startswith('- '):
                            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                            apply_font_to_paragraph(p, BRAND_FONT)
                        elif line.strip().startswith('* '):
                            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                            apply_font_to_paragraph(p, BRAND_FONT)
                        elif re.match(r'^\d+\.\s', line.strip()):
                            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
                            apply_font_to_paragraph(p, BRAND_FONT)
                        else:
                            p = doc.add_paragraph(line.strip())
                            apply_font_to_paragraph(p, BRAND_FONT)
    
    # Save
    output = Path(output_path) if output_path else folder / 'risk_analysis_report.docx'
    doc.save(str(output))
    print(f'Created: {output}')
    return output


def generate_test_coverage_docx(folder_path, output_path=None):
    """
    Generate a formatted DOCX from test_coverage_map.md.
    
    Args:
        folder_path: Path to the risk analysis folder
        output_path: Optional custom output path
    
    Returns:
        Path to the generated DOCX file
    """
    folder = Path(folder_path)
    md_path = folder / 'test_coverage_map.md'
    
    if not md_path.exists():
        print(f"Test coverage map not found: {md_path} - skipping")
        return None
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    set_document_font(doc, BRAND_FONT)
    
    # Title
    title = doc.add_heading('Test Case Coverage Map', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PEPPERCORN
        run.font.name = BRAND_FONT
    
    # Info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Generated: ').bold = True
    info.add_run(f'{datetime.now().strftime("%Y-%m-%d")}\n')
    info.add_run('Source: ').bold = True
    info.add_run(str(md_path))
    apply_font_to_paragraph(info, BRAND_FONT)
    
    # Parse sections and tables similar to risk analysis
    # This is a simplified version - extend as needed
    
    lines = content.split('\n')
    current_section = None
    table_lines = []
    
    for line in lines:
        if line.startswith('## '):
            # Flush table
            if table_lines:
                headers, rows = parse_markdown_table(table_lines)
                if headers and rows:
                    add_table_with_header(doc, headers, rows)
                table_lines = []
            
            add_styled_heading(doc, line[3:].strip(), 2)
            current_section = line[3:].strip()
        
        elif line.startswith('### '):
            # Flush table
            if table_lines:
                headers, rows = parse_markdown_table(table_lines)
                if headers and rows:
                    add_table_with_header(doc, headers, rows)
                table_lines = []
            
            add_styled_heading(doc, line[4:].strip(), 3)
        
        elif line.strip().startswith('|'):
            table_lines.append(line)
        
        elif table_lines and not line.strip():
            # End of table
            headers, rows = parse_markdown_table(table_lines)
            if headers and rows:
                add_table_with_header(doc, headers, rows)
                doc.add_paragraph()
            table_lines = []
        
        elif line.strip() and not line.startswith('#'):
            if line.strip().startswith('- '):
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                apply_font_to_paragraph(p, BRAND_FONT)
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.strip().strip('*')).bold = True
                apply_font_to_paragraph(p, BRAND_FONT)
            elif not line.startswith('|'):
                p = doc.add_paragraph(line.strip())
                apply_font_to_paragraph(p, BRAND_FONT)
    
    # Flush remaining table
    if table_lines:
        headers, rows = parse_markdown_table(table_lines)
        if headers and rows:
            add_table_with_header(doc, headers, rows)
    
    # Save
    output = Path(output_path) if output_path else folder / 'test_coverage_map.docx'
    doc.save(str(output))
    print(f'Created: {output}')
    return output


def detect_folder_type(folder_path):
    """
    Detect whether this is a risk analysis folder or test jam folder.
    
    Returns:
        tuple: (folder_type, available_files)
        folder_type: 'risk_analysis', 'test_jam', or 'unknown'
    """
    folder = Path(folder_path)
    available_files = []
    
    # Check for risk analysis files
    if (folder / 'risk_analysis_report.md').exists():
        available_files.append('risk_analysis_report.md')
    
    # Check for test coverage map
    if (folder / 'test_coverage_map.md').exists():
        available_files.append('test_coverage_map.md')
    
    # Check for test jam files
    if (folder / 'testjam_all_test_cases.csv').exists():
        available_files.append('testjam_all_test_cases.csv')
    
    # Determine folder type
    if 'risk_analysis_report.md' in available_files:
        return 'risk_analysis', available_files
    elif 'testjam_all_test_cases.csv' in available_files:
        return 'test_jam', available_files
    elif 'test_coverage_map.md' in available_files:
        return 'coverage_only', available_files
    else:
        return 'unknown', available_files


def generate_all_docx(folder_path):
    """Generate all DOCX files for a QualityForge folder (risk analysis or test jam)."""
    folder_type, available_files = detect_folder_type(folder_path)
    
    print(f"\n📄 Generating DOCX files with company branding...")
    print(f"   Font: {BRAND_FONT}")
    print(f"   Folder: {folder_path}")
    print(f"   Type: {folder_type}")
    print(f"   Available: {', '.join(available_files)}\n")
    
    outputs = []
    
    # Generate risk analysis report if it exists
    if 'risk_analysis_report.md' in available_files:
        try:
            risk_docx = generate_risk_analysis_docx(folder_path)
            outputs.append(risk_docx)
        except FileNotFoundError as e:
            print(f"Warning: {e}")
    
    # Generate test coverage map if it exists
    if 'test_coverage_map.md' in available_files:
        try:
            coverage_docx = generate_test_coverage_docx(folder_path)
            if coverage_docx:
                outputs.append(coverage_docx)
        except FileNotFoundError as e:
            print(f"Warning: {e}")
    
    if not outputs:
        print("⚠️  No markdown files found to convert.")
        print("   Expected: risk_analysis_report.md and/or test_coverage_map.md")
        return outputs
    
    print(f"\n✅ Generated {len(outputs)} DOCX file(s)")
    return outputs


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Generate formatted DOCX files from QualityForge Risk Analysis outputs'
    )
    parser.add_argument(
        'folder',
        help='Path to the risk analysis folder containing markdown files'
    )
    parser.add_argument(
        '--risk-only',
        action='store_true',
        help='Only generate the risk analysis report'
    )
    parser.add_argument(
        '--coverage-only',
        action='store_true',
        help='Only generate the test coverage map'
    )
    
    args = parser.parse_args()
    
    if not os.path.isdir(args.folder):
        print(f"Error: Folder not found: {args.folder}")
        sys.exit(1)
    
    if args.risk_only:
        generate_risk_analysis_docx(args.folder)
    elif args.coverage_only:
        generate_test_coverage_docx(args.folder)
    else:
        generate_all_docx(args.folder)


if __name__ == '__main__':
    main()
